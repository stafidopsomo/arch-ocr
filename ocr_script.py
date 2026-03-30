"""
OCR → DOCX Tool for Greek Property Documents (HTK)
===================================================
Converts one or more scanned PDF documents into a single Word document.
- Page 1: HTK cross-reference summary table (fields found across all documents)
- Remaining pages: per-page table with [image | OCR text]

Usage:
    python ocr_script.py <pdf1> [pdf2] [pdf3] ...  [--output output/result.docx]
    python ocr_script.py /path/to/folder/           [--output output/result.docx]
"""

import argparse
import hashlib
import json
import os
import re
import sys
import tempfile
from collections import defaultdict
from pathlib import Path

import cv2
import numpy as np
import pytesseract
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pdf2image import convert_from_path
from PIL import Image

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Adjust this if Tesseract is not on PATH (Windows common location)
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Adjust this if poppler is not on PATH (Windows)
POPPLER_PATH = None  # e.g. r"C:\poppler\bin"

OCR_LANG = "ell+grc"   # Modern Greek + ancient/polytonic Greek
OCR_CONFIG = "--psm 6 --oem 1"

# HTK fields: (display_name, list_of_regex_patterns_to_search)
HTK_FIELDS = [
    ("Όνομα",           [r"(?:Ονομα|ΟΝΟΜΑ|ΠΑΥΛΟΣ|Όνομα)[:\s]+([Α-Ωα-ωΆ-Ώ\w]+)"]),
    ("Επώνυμο",         [r"(?:Επώνυμο|ΕΠΩΝΥΜΟ|υπογεγραμμένος)[:\s]+.*?([Α-Ωα-ω]{3,})"]),
    ("Πατρώνυμο",       [r"(?:πατρός|ΠΑΤΡΩΝΥΜΟ|Πατρώνυμο|πατρώνυμ)[:\s]+([Α-Ωα-ωΆ-Ώ\w]+)"]),
    ("Μητρώνυμο",       [r"(?:μητρός|ΜΗΤΡΩΝΥΜΟ|μητρώνυμ)[:\s]+([Α-Ωα-ωΆ-Ώ\w]+)"]),
    ("Έτος γέννησης",   [r"(?:γεννηθείς|γεννηθεί|γεννημένος|κατοικ)[^\d]*(\d{4})"]),
    ("Τόπος γέννησης",  [r"(?:γεννηθείς|γεννηθεί)\s+(?:εν|εις|στ[ηο])\s+([Α-Ωα-ωΆ-Ώ\w]+)"]),
    ("Διεύθυνση",       [r"(?:οδός|οδό|επί της οδού|οδ\.)\s*([Α-Ωα-ωΆ-Ώ\w\s\d]+?)(?:\s*αρι|,|\.|\d)",
                          r"(?:ΟΔΟΣ|ΟΔΟ)[:\s]+([Α-Ωα-ωΆ-ΏA-Za-z\s\d]+?)(?:\n|,)"]),
    ("Αριθμός ακινήτου",[r"(?:αριθ\.|αριθμ\.|αρ\.)\s*(\d+[\w\-]*)"]),
    ("Χρήση ακινήτου",  [r"(?:Χρήσις|Χρήση)\s+(?:ακινήτου)?\s*[:\.]?\s*([Α-Ωα-ωΆ-ΏA-Za-z\s]+?)(?:\n|\()"]),
    ("Αριθμός ορόφων",  [r"(?:Αριθμ[οό]ς\s+ορόφων|ορόφων\s+κα[ιί]\s+στάσεων)[^\d]*([0-9\-ΗΜΙημι\s]+)"]),
    ("ΑΔΤ",             [r"(?:δελτίου\s+ταυτότητος|ταυτότητας|ΑΔΤ|Αρ\.\s*Δελτ)[^\w]*([Α-ΩA-Z]{1,2}\s*\d{5,9})"]),
    ("Τηλέφωνο",        [r"(?:ΤΗΛΕΦ|Τηλέφ|τηλ|ΤΗΛ)[.\s:]*(\d[\d\s\-/]{6,14})"]),
    ("ΑΦΜ",             [r"(?:ΑΦΜ|Α\.Φ\.Μ|αριθμ\.\s*φορ)[.\s:]*(\d{9})"]),
]

# Flatten HTK_FIELDS so values are always lists (handle tuple vs list inconsistency)
HTK_FIELDS = [
    (name, patterns if isinstance(patterns, list) else [patterns])
    for name, patterns in HTK_FIELDS
]

# Colors
HEADER_BG = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
HEADER_BG_HEX = "1F497D"
ALT_ROW_BG = RGBColor(0xD9, 0xE1, 0xF2)  # light blue
ALT_ROW_BG_HEX = "D9E1F2"

# ---------------------------------------------------------------------------
# Image preprocessing
# ---------------------------------------------------------------------------

def preprocess_image(pil_img: Image.Image) -> Image.Image:
    """Enhance image quality for OCR: grayscale, denoise, deskew, threshold."""
    img = np.array(pil_img.convert("RGB"))
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

    # Upscale if image is small (improves OCR on low-res scans)
    h, w = gray.shape
    if w < 1500:
        scale = 1500 / w
        gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    # CLAHE for contrast
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # Denoise
    gray = cv2.fastNlMeansDenoising(gray, h=10)

    # Deskew
    gray = _deskew(gray)

    # Otsu threshold
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    return Image.fromarray(binary)


def _deskew(gray: np.ndarray) -> np.ndarray:
    """Detect and correct rotation using Hough lines."""
    try:
        edges = cv2.Canny(gray, 50, 150, apertureSize=3)
        lines = cv2.HoughLinesP(edges, 1, np.pi / 180, threshold=100,
                                 minLineLength=100, maxLineGap=10)
        if lines is None:
            return gray
        angles = []
        for line in lines:
            x1, y1, x2, y2 = line[0]
            if x2 - x1 != 0:
                angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
                if -15 < angle < 15:
                    angles.append(angle)
        if not angles:
            return gray
        median_angle = float(np.median(angles))
        if abs(median_angle) < 0.3:
            return gray
        h, w = gray.shape
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, median_angle, 1.0)
        rotated = cv2.warpAffine(gray, M, (w, h),
                                  flags=cv2.INTER_CUBIC,
                                  borderMode=cv2.BORDER_REPLICATE)
        return rotated
    except Exception:
        return gray


# ---------------------------------------------------------------------------
# OCR
# ---------------------------------------------------------------------------

def run_ocr(pil_img: Image.Image) -> str:
    """Run Tesseract OCR on a PIL image, returns extracted text."""
    preprocessed = preprocess_image(pil_img)
    text = pytesseract.image_to_string(preprocessed, lang=OCR_LANG, config=OCR_CONFIG)
    return text


# ---------------------------------------------------------------------------
# HTK field extraction
# ---------------------------------------------------------------------------

def extract_htk_fields(all_page_texts: list[tuple[str, int, str]]) -> dict:
    """
    Scan OCR text from all pages and extract HTK fields.

    all_page_texts: list of (pdf_name, page_number, ocr_text)

    Returns: dict of {field_name: {"value": str, "occurrences": [(pdf, page)]}}
    """
    results = {}
    for field_name, patterns in HTK_FIELDS:
        occurrences = []
        best_value = None
        value_counts = defaultdict(int)

        for pdf_name, page_num, text in all_page_texts:
            for pattern in patterns:
                matches = re.findall(pattern, text, re.IGNORECASE | re.UNICODE)
                for m in matches:
                    val = m.strip() if isinstance(m, str) else " ".join(m).strip()
                    val = re.sub(r"\s+", " ", val).strip()
                    if len(val) > 1:
                        occurrences.append((pdf_name, page_num))
                        value_counts[val] += 1

        if value_counts:
            best_value = max(value_counts, key=value_counts.get)

        # Deduplicate occurrences (keep unique page references)
        seen = set()
        unique_occ = []
        for occ in occurrences:
            if occ not in seen:
                seen.add(occ)
                unique_occ.append(occ)

        results[field_name] = {
            "value": best_value or "—",
            "occurrences": unique_occ,
        }
    return results


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, rgb: RGBColor):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    color_hex = rgb if isinstance(rgb, str) else f"{rgb.red:02X}{rgb.green:02X}{rgb.blue:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_border(cell, border_size=4):
    """Add thin borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para_style(para, bold=False, size=10, color=None, align=None):
    run = para.runs[0] if para.runs else para.add_run()
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align


def add_htk_summary_page(doc: Document, htk_results: dict, pdf_names: list[str]):
    """Add the first page: HTK cross-reference summary table."""
    heading = doc.add_heading("ΣΤΟΙΧΕΙΑ ΓΙΑ ΗΛΕΚΤΡΟΝΙΚΗ ΤΑΥΤΟΤΗΤΑ ΚΤΙΡΙΟΥ (ΗΤΚ)", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph(f"Επεξεργάστηκαν {len(pdf_names)} έγγραφο/α: {', '.join(pdf_names)}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # Table: Πεδίο | Καλύτερη τιμή | Βρέθηκε σε (σελίδες) | Σύνολο αναφορών
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["Πεδίο ΗΤΚ", "Καλύτερη τιμή", "Βρέθηκε σε (σελίδα/έγγραφο)", "Σύνολο αναφορών"]
    for i, (cell, text) in enumerate(zip(hdr_cells, headers)):
        _set_cell_bg(cell, HEADER_BG_HEX)
        _set_cell_border(cell)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, (field_name, _) in enumerate(HTK_FIELDS):
        data = htk_results.get(field_name, {"value": "—", "occurrences": []})
        occs = data["occurrences"]

        # Format occurrence list
        if occs:
            occ_parts = []
            for pdf_name, page_num in occs:
                occ_parts.append(f"Σελ.{page_num} ({Path(pdf_name).stem})")
            occ_str = ", ".join(occ_parts)
        else:
            occ_str = "Δεν βρέθηκε"

        row_cells = table.add_row().cells
        row_data = [field_name, data["value"], occ_str, str(len(occs))]

        for cell, text in zip(row_cells, row_data):
            _set_cell_border(cell)
            if row_idx % 2 == 0:
                _set_cell_bg(cell, ALT_ROW_BG_HEX)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)

    # Set column widths
    widths = [Inches(1.5), Inches(2.0), Inches(3.0), Inches(0.8)]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    doc.add_page_break()


def add_page_spread(doc: Document, pdf_name: str, page_num: int,
                     pil_img: Image.Image, ocr_text: str, img_dir: str):
    """Add a two-column table row: [page image] | [OCR text]."""
    # Section label
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = label.add_run(f"📄 {Path(pdf_name).name}  —  Σελίδα {page_num}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Save original page image for embedding
    img_path = os.path.join(img_dir, f"page_{Path(pdf_name).stem}_{page_num}.png")
    pil_img.save(img_path)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    row = table.rows[0]

    # Left cell: image
    left_cell = row.cells[0]
    left_cell.width = Inches(3.5)
    _set_cell_border(left_cell)
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = lp.add_run()
    run_img.add_picture(img_path, width=Inches(3.3))

    # Right cell: OCR text
    right_cell = row.cells[1]
    right_cell.width = Inches(4.5)
    _set_cell_border(right_cell)
    rp = right_cell.paragraphs[0]
    rp.clear()
    run_txt = rp.add_run(ocr_text.strip() if ocr_text.strip() else "(Δεν αναγνωρίστηκε κείμενο)")
    run_txt.font.size = Pt(8)
    run_txt.font.name = "Calibri"

    doc.add_paragraph()  # spacing between pages


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# OCR cache (saves results to avoid re-running OCR on unchanged PDFs)
# ---------------------------------------------------------------------------

CACHE_DIR = Path("ocr_cache")

def _pdf_hash(pdf_path: str) -> str:
    h = hashlib.md5()
    with open(pdf_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()

def load_cache(pdf_path: str) -> list | None:
    """Return cached list of (page_num, text) or None if not cached."""
    CACHE_DIR.mkdir(exist_ok=True)
    key = _pdf_hash(pdf_path)
    cache_file = CACHE_DIR / f"{key}.json"
    if cache_file.exists():
        with open(cache_file) as f:
            return json.load(f)
    return None

def save_cache(pdf_path: str, pages: list):
    """Save list of (page_num, text) to cache."""
    CACHE_DIR.mkdir(exist_ok=True)
    key = _pdf_hash(pdf_path)
    cache_file = CACHE_DIR / f"{key}.json"
    with open(cache_file, "w", encoding="utf-8") as f:
        json.dump(pages, f, ensure_ascii=False, indent=2)


def collect_pdfs(inputs: list[str]) -> list[str]:
    """Resolve input paths (files or folders) to a list of PDF paths."""
    pdfs = []
    for inp in inputs:
        p = Path(inp)
        if p.is_dir():
            pdfs.extend(sorted(p.glob("*.pdf")))
            pdfs.extend(sorted(p.glob("*.PDF")))
        elif p.is_file() and p.suffix.lower() == ".pdf":
            pdfs.append(p)
        else:
            print(f"[WARN] Skipping '{inp}' — not a PDF file or directory.")
    return [str(p) for p in pdfs]


def main():
    parser = argparse.ArgumentParser(
        description="OCR Greek property PDFs → DOCX with HTK summary"
    )
    parser.add_argument(
        "inputs", nargs="+",
        help="One or more PDF file paths, or a folder containing PDFs"
    )
    parser.add_argument(
        "--output", default="output/result.docx",
        help="Output DOCX path (default: output/result.docx)"
    )
    parser.add_argument(
        "--poppler-path", default=None,
        help="Path to poppler bin directory (Windows, if not on PATH)"
    )
    parser.add_argument(
        "--tesseract-path", default=None,
        help="Path to tesseract.exe (Windows, if not on PATH)"
    )
    args = parser.parse_args()

    if args.tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = args.tesseract_path

    poppler_path = args.poppler_path or POPPLER_PATH

    # Resolve PDFs
    pdf_paths = collect_pdfs(args.inputs)
    if not pdf_paths:
        print("[ERROR] No PDF files found. Exiting.")
        sys.exit(1)

    print(f"[INFO] Processing {len(pdf_paths)} PDF(s):")
    for p in pdf_paths:
        print(f"       {p}")

    # Ensure output directory exists
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set page margins (narrow for more content space)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    all_page_texts = []   # (pdf_path, page_num, ocr_text)
    all_pages = []        # (pdf_path, page_num, pil_img, ocr_text)

    with tempfile.TemporaryDirectory() as tmp_dir:
        # Phase 1: convert & OCR all pages (with cache)
        for pdf_path in pdf_paths:
            pdf_name = Path(pdf_path).name
            cached = load_cache(pdf_path)

            print(f"\n[INFO] Converting '{pdf_name}' to images...")
            try:
                pil_pages = convert_from_path(
                    pdf_path, dpi=300,
                    poppler_path=poppler_path
                )
            except Exception as e:
                print(f"[ERROR] Could not convert '{pdf_name}': {e}")
                continue

            if cached and len(cached) == len(pil_pages):
                print(f"  [CACHE] Using cached OCR for '{pdf_name}' ({len(pil_pages)} pages)")
                for page_num, pil_img in enumerate(pil_pages, start=1):
                    text = cached[page_num - 1][1]
                    all_page_texts.append((pdf_path, page_num, text))
                    all_pages.append((pdf_path, page_num, pil_img, text))
            else:
                page_cache = []
                for page_num, pil_img in enumerate(pil_pages, start=1):
                    print(f"  OCR page {page_num}/{len(pil_pages)}...", end=" ", flush=True)
                    try:
                        text = run_ocr(pil_img)
                        print("OK")
                    except Exception as e:
                        text = f"[OCR error: {e}]"
                        print(f"ERROR: {e}")
                    page_cache.append((page_num, text))
                    all_page_texts.append((pdf_path, page_num, text))
                    all_pages.append((pdf_path, page_num, pil_img, text))
                save_cache(pdf_path, page_cache)
                print(f"  [CACHE] Saved OCR cache for '{pdf_name}'")

        # Phase 2: extract HTK fields
        print("\n[INFO] Extracting HTK fields...")
        htk_results = extract_htk_fields(all_page_texts)

        pdf_names = [Path(p).name for p in pdf_paths]

        # Phase 3: build DOCX
        print("[INFO] Building DOCX...")

        # Page 1: HTK summary
        add_htk_summary_page(doc, htk_results, pdf_names)

        # Remaining pages: image + OCR text
        for pdf_path, page_num, pil_img, ocr_text in all_pages:
            add_page_spread(doc, pdf_path, page_num, pil_img, ocr_text, tmp_dir)

        # Save
        doc.save(str(out_path))
        print(f"\n[DONE] Saved: {out_path.resolve()}")


if __name__ == "__main__":
    main()
